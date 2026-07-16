#!/usr/bin/env python3
"""
Generates the FRAD Foundation Headless WordPress CMS Administrator & Content Guide (.docx)
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_with_spacing(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14 if level == 1 else (10 if level == 2 else 6))
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    return h

def add_callout(doc, title, text, border_color="115E3B", bg_color="F0F7F4"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    # Border left thick
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left_b = OxmlElement('w:left')
    left_b.set(qn('w:val'), 'single')
    left_b.set(qn('w:sz'), '24') # 3pt
    left_b.set(qn('w:space'), '0')
    left_b.set(qn('w:color'), border_color)
    tcBorders.append(left_b)
    for b in ['top', 'bottom', 'right']:
        node = OxmlElement(f'w:{b}')
        node.set(qn('w:val'), 'none')
        tcBorders.append(node)
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"{title.upper()}\n")
    run_t.bold = True
    run_t.font.color.rgb = RGBColor(17, 94, 59)
    run_t.font.size = Pt(10)
    
    run_b = p.add_run(text)
    run_b.font.size = Pt(10.5)
    run_b.font.color.rgb = RGBColor(30, 41, 59)
    doc.add_paragraph() # Spacer

def build_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header Row
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], "115E3B")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=140, right=140)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
            
    # Data Rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=140, right=140)
            p = row_cells[c_idx].paragraphs[0]
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(15, 23, 42)
    doc.add_paragraph()
    return table

def main():
    doc = docx.Document()
    
    # Page setup
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
        
    # Title / Header
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run("FRAD Foundation — Headless WordPress CMS Administrator & Content Guide")
    r_title.bold = True
    r_title.font.size = Pt(22)
    r_title.font.color.rgb = RGBColor(17, 94, 59) # FRAD Green
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(18)
    r_sub = p_sub.add_run("Comprehensive Technical Handbook for Managing Every Website Section via WordPress Backend")
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = RGBColor(71, 85, 105)
    
    add_callout(doc, "Executive Summary & Architectural Guarantee",
                "Every single frontend section across the FRAD Foundation Next.js web application is wired to accept dynamic content from your headless WordPress backend via WPGraphQL. When configured in WordPress, your edits immediately control website headings, visual features, emergency alerts, programme areas, partner logos, reports, and global footer settings. If WordPress is offline or a slot is left empty, the site automatically falls back to curated static editorial content, ensuring zero visual breakage.")

    # Section 1: Overview of Controlled Sections
    add_heading_with_spacing(doc, "1. Sitewide Inventory: What is Controlled by WordPress", level=1)
    
    doc.add_paragraph(
        "The FRAD Foundation application divides content into three main management models in WordPress:\n"
        "1. Custom Post Types (CPTs): Structured records such as Emergency Alerts, Projects, Stories, Sectors, Partners, Reports, and Media Assets.\n"
        "2. Editorial Content Slots (fradContentSlots): Flexible text, image, video, exhibit metadata, and CTA button overrides for specific page sections (Heroes, Visual Features, Call to Action banners, Dossiers).\n"
        "3. Global Site Settings (ACF Options Page): Organization-wide details such as CAC Registration Number, office addresses, contact emails, and hotlines."
    )
    
    headers = ["Frontend Section / Component", "WordPress Data Source", "Key Editable Fields", "Fallback Behavior"]
    rows = [
        ["Emergency Banner (Sitewide)", "fradAlerts (Custom Post Type)", "Title, Alert Message, Priority, Active Flag, CTA Button Link & Label", "Displays curated static alerts or hides cleanly"],
        ["Hero Banner (Homepage / Sections)", "fradContentSlots (Key: e.g., home.hero)", "Headline, Subtext, Background Image/Video, Primary/Secondary/Tertiary CTAs", "Displays static field operations hero"],
        ["Visual Feature Frames", "fradContentSlots (Key: e.g., home.where_we_work.visual)", "Image, Video Reel URL, Duration, Eyebrow, Title, Caption, Exhibit Metadata (Location, Subject, Handling Notes, Stamps)", "Displays default editorial field photo/video frame"],
        ["Call To Action Section", "fradContentSlots (Key: e.g., home.cta)", "Eyebrow, Headline, Subtext, Primary & Secondary CTA Buttons", "Displays default 'Work with FRAD' partnership banner"],
        ["Evidence Dossier Section", "fradContentSlots (Key: e.g., home.dossier)", "Eyebrow, Headline, Subtext, Repeater Rows (Label & Value pair), Action Link", "Displays static verified evidence dossier"],
        ["Sector Explorer Dashboard", "fradSectors (Custom Post Type)", "Sector Name, Slug, Icon, Short Description, Full Note, Key Activities", "Displays static 6 core FRAD programme sectors"],
        ["Accountability Meter", "fradContentSlots (Key: about.accountability.meter)", "Headline, Subtext, Metric Repeater Rows", "Displays default accountability & safeguarding metrics"],
        ["Featured Projects Grid", "fradProjects (Custom Post Type)", "Project Title, Sector, Location, Summary, Donor, Status, Featured Flag", "Displays static featured humanitarian projects"],
        ["Featured Stories Grid", "fradStories (Custom Post Type)", "Story Title, Excerpt, Content, Featured Image, Author, Featured Flag", "Displays static community impact stories"],
        ["Partner Proof Strip", "fradPartners (Custom Post Type)", "Partner Name, Logo Image, Website URL, Approved to Display Flag", "Displays approved institutional partners or empty state"],
        ["Report Shelf Grid", "fradReports (Custom Post Type)", "Report Title, Year, Type, Sector, Summary, PDF Download URL", "Displays static accountability reports & audits"],
        ["Media Showcase & Library", "fradMediaAssets (Custom Post Type)", "Asset Title, Image/Video URL, Sector, Location, Consent Filed Flag, Public Safe Flag", "Displays curated field photography showcase"],
        ["Sitewide Footer & Legal Line", "ACF Options: frad_site_settings", "CAC Registration Number, Office Locations, Contact Email & Phone", "Displays default 'CAC/IT/NO/139393' and static footer"]
    ]
    build_table(doc, headers, rows)

    # Section 2: Step-by-Step Configuration Guide for Custom Post Types
    add_heading_with_spacing(doc, "2. Configuring Custom Post Types (CPTs) in WordPress", level=1)
    
    doc.add_paragraph(
        "To manage structured items, install Advanced Custom Fields (ACF Pro) and WPGraphQL on your WordPress backend. Create the following Custom Post Types and ensure 'Show in GraphQL' is enabled for each."
    )

    add_heading_with_spacing(doc, "A. Emergency Alerts (GraphQL Single Name: emergencyAlert, Plural: fradAlerts)", level=2)
    doc.add_paragraph("Go to WordPress Admin > Emergency Alerts > Add New. Create an ACF Field Group called 'Alert Fields' mapped to Post Type = Emergency Alert:")
    build_table(doc, [
        "ACF Field Label", "Field Name (GraphQL)", "Field Type", "Instructions & Options"
    ], [
        ["Alert Message", "message", "Textarea", "Full description of the emergency or operational alert."],
        ["CTA Button Text", "ctaText", "Text", "e.g., 'View Situation Report' or 'Support Response'"],
        ["CTA Button Link", "ctaLink", "URL / Text", "Relative path e.g. '/reports/sitrep-04' or external URL"],
        ["Priority Tone", "priority", "Select", "Choices: low, medium, high, critical"],
        ["Active Alert Flag", "active", "True / False", "Set to True to broadcast sitewide immediately."]
    ])

    add_heading_with_spacing(doc, "B. Programme Sectors (GraphQL Plural: fradSectors)", level=2)
    doc.add_paragraph("Go to WordPress Admin > Sectors > Add New. Create an ACF Field Group called 'Sector Fields' mapped to Post Type = Sector:")
    build_table(doc, [
        "ACF Field Label", "Field Name (GraphQL)", "Field Type", "Instructions & Options"
    ], [
        ["URL Slug", "slug", "Text", "e.g. 'nutrition-health' or 'wash'"],
        ["Short Description", "shortDescription", "Textarea", "1-2 sentence overview displayed on dashboard cards."],
        ["Full Description", "fullDescription", "Textarea", "Expanded operational note for detailed sector pages."],
        ["Icon Identifier", "icon", "Text", "Lucide/Custom icon key e.g. 'Heart', 'Droplets', 'Shield'"],
        ["Key Activities", "keyActivities", "Repeater / Text lines", "List of technical interventions delivered in this sector."]
    ])

    add_heading_with_spacing(doc, "C. Institutional Partners (GraphQL Plural: fradPartners)", level=2)
    doc.add_paragraph("Go to WordPress Admin > Partners > Add New. Create an ACF Field Group called 'Partner Fields':")
    build_table(doc, [
        "ACF Field Label", "Field Name (GraphQL)", "Field Type", "Instructions & Options"
    ], [
        ["Partner Logo", "logo", "Image", "High-resolution transparent PNG logo."],
        ["Website URL", "website", "URL", "Official institutional website link."],
        ["Approved for Display", "approvedToDisplay", "True / False", "Must be True for logo to appear on public Partner Strip."]
    ])

    # Section 3: Step-by-Step Configuration Guide for Editorial Content Slots
    add_heading_with_spacing(doc, "3. Managing Page Sections via Content Slots (fradContentSlots)", level=1)
    
    doc.add_paragraph(
        "Content Slots allow you to override specific headings, visual frames, and CTA sections across any page without touching code. In WordPress Admin, create a Custom Post Type named 'Content Slot' (GraphQL Plural: fradContentSlots)."
    )
    
    add_callout(doc, "Crucial Step: Naming Your Content Slot Slug",
                "When creating a Content Slot post in WordPress, you MUST set the Post Slug (permalink slug) to match the exact Slot Key listed below (e.g., home.cta or home.where_we_work.visual). The Next.js frontend fetches the slot matching that slug.")

    add_heading_with_spacing(doc, "Master Slot Key Registry", level=2)
    build_table(doc, [
        "WordPress Post Slug (Slot Key)", "Page / Section Location", "What It Controls on the Website"
    ], [
        ["home.hero", "Homepage Top Banner", "Main display title, subtext, background image/video, and hero CTA buttons."],
        ["home.where_we_work.visual", "Homepage Section 02 Visual", "Water access photo/video reel, caption, and Exhibit Metadata (Location, Subject, Stamps)."],
        ["home.what_we_do.visual", "Homepage Section 03 Visual", "Programme outreach reel, duration badge, and video playback frame."],
        ["home.cta", "Homepage Section 05 CTA Banner", "Work With FRAD partnership banner headline, subtext, and primary/secondary buttons."],
        ["about.accountability.meter", "Accountability & Governance Page", "Headline, description, and dynamic accountability indicator rows."]
    ])

    add_heading_with_spacing(doc, "ACF Field Group Schema for Content Slots ('slotFields')", level=2)
    build_table(doc, [
        "ACF Field Label", "Field Name (GraphQL)", "Field Type", "Description / Usage"
    ], [
        ["Eyebrow Kicker", "eyebrow", "Text", "Small uppercase kicker above the main title."],
        ["Headline", "headline", "Text", "Primary section title."],
        ["Subtext / Description", "subtext", "Textarea", "Supporting explanatory text."],
        ["Background Image", "image", "Image", "Main feature image or hero poster."],
        ["Video URL", "videoUrl", "URL / Text", "MP4 background video URL."],
        ["Primary CTA Text", "ctaText", "Text", "Text for main action button."],
        ["Primary CTA Link", "ctaLink", "Text", "Destination path for main action button."],
        ["Secondary CTA Text", "secondaryCtaText", "Text", "Text for outline secondary button."],
        ["Secondary CTA Link", "secondaryCtaLink", "Text", "Destination path for secondary button."],
        ["Exhibit Metadata Group", "exhibit", "Group", "Contains location (Text), subject (Text), handling (Text), and stamps (Repeater of label+tone)."]
    ])

    # Section 4: Global Footer & Site Settings
    add_heading_with_spacing(doc, "4. Global Footer & Site Settings (ACF Options Page)", level=1)
    doc.add_paragraph(
        "To control global legal and organization data displayed in the Footer:\n"
        "1. Create an ACF Options Page in WordPress titled 'FRAD Site Settings' (GraphQL Option Name: fradSiteSettings).\n"
        "2. Add a field named cacRegistrationNumber (Text). When populated (e.g., 'CAC/IT/NO/139393' or updated official registration), it immediately updates the legal copyright line at the bottom of every page.\n"
        "3. Add contactEmail and phoneNumbers to control sitewide contact blocks."
    )

    # Section 5: Verification & Cache Invalidation
    add_heading_with_spacing(doc, "5. Verification & Cache Invalidation Workflow", level=1)
    doc.add_paragraph(
        "The FRAD Next.js frontend uses intelligent revalidation (ISR) to keep pages lightning-fast while staying updated:\n"
        "• Editorial Content Slots revalidate every 120 seconds automatically.\n"
        "• Emergency Alerts revalidate every 60 seconds automatically.\n"
        "• To trigger instantaneous updates on publication, install the 'WPGraphQL Smart Cache' or configure an on-save Webhook pointing to your Next.js deployment URL."
    )

    doc.save("/Users/macbookpro/Documents/fradwebsite/FRAD_WordPress_CMS_Administrator_Guide.docx")
    print("Successfully generated Word document at /Users/macbookpro/Documents/fradwebsite/FRAD_WordPress_CMS_Administrator_Guide.docx")

if __name__ == '__main__':
    main()
